from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from llm.openai_client import embeddings
from langchain_community.vectorstores import Chroma

import tempfile
import os


def add_documents(filename, file_bytes):

    # Save file temporarily
    temp_file = tempfile.NamedTemporaryFile(delete=False)
    temp_file.write(file_bytes)
    temp_file.close()

    text = ""

    # Load depending on file type
    if filename.endswith(".txt"):

        with open(temp_file.name, "r", encoding="utf-8") as f:
            text = f.read()

    elif filename.endswith(".pdf"):

        from pypdf import PdfReader
        reader = PdfReader(temp_file.name)

        for page in reader.pages:
            text += page.extract_text()

    elif filename.endswith(".docx"):

        import docx
        doc = docx.Document(temp_file.name)

        for para in doc.paragraphs:
            text += para.text + "\n"

    else:
        return "Unsupported file type"

    os.unlink(temp_file.name)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100
    )

    docs = splitter.split_documents([Document(page_content=text)])

    # Check if vector store exists, if so load it and add documents, otherwise create new
    if os.path.exists("vector_store"):
        vectorstore = Chroma.load_local(
            "vector_store",
            embeddings,
            allow_dangerous_deserialization=True
        )
        vectorstore.add_documents(docs)
    else:
        vectorstore = Chroma.from_documents(docs, embeddings)

    vectorstore.save_local("vector_store")

    return "File indexed successfully"